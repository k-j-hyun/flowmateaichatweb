"""
Word 문서 생성 모듈
마크다운 텍스트를 전문적인 DOCX로 변환
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import os
import logging
from datetime import datetime
from io import BytesIO

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DOCXGenerator:
    """Word 문서 생성기 클래스"""
    
    def __init__(self):
        """초기화"""
        self.doc = None
        self.theme_colors = {
            'primary': RGBColor(99, 102, 241),    # 보라색
            'secondary': RGBColor(139, 92, 246),  # 연한 보라색
            'accent': RGBColor(124, 58, 237),     # 진한 보라색
            'text': RGBColor(74, 85, 104),        # 회색
            'light_gray': RGBColor(128, 128, 128) # 연한 회색
        }
    
    def create_document(self):
        """새 문서 생성"""
        try:
            self.doc = Document()
            self._create_styles()
            logger.info("문서 객체 생성 완료")
            return True
        except Exception as e:
            logger.error(f"문서 생성 실패: {e}")
            return False
    
    def _create_styles(self):
        """커스텀 스타일 생성"""
        try:
            styles = self.doc.styles
            
            # 제목 1 스타일
            if 'CustomHeading1' not in [s.name for s in styles]:
                h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
                h1_font = h1_style.font
                h1_font.name = '맑은 고딕'
                h1_font.size = Pt(24)
                h1_font.bold = True
                h1_font.color.rgb = self.theme_colors['primary']
                h1_style.paragraph_format.space_after = Pt(12)
                h1_style.paragraph_format.space_before = Pt(24)
            
            # 제목 2 스타일
            if 'CustomHeading2' not in [s.name for s in styles]:
                h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
                h2_font = h2_style.font
                h2_font.name = '맑은 고딕'
                h2_font.size = Pt(18)
                h2_font.bold = True
                h2_font.color.rgb = self.theme_colors['secondary']
                h2_style.paragraph_format.space_after = Pt(10)
                h2_style.paragraph_format.space_before = Pt(16)
            
            # 제목 3 스타일
            if 'CustomHeading3' not in [s.name for s in styles]:
                h3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
                h3_font = h3_style.font
                h3_font.name = '맑은 고딕'
                h3_font.size = Pt(14)
                h3_font.bold = True
                h3_font.color.rgb = self.theme_colors['accent']
                h3_style.paragraph_format.space_after = Pt(8)
                h3_style.paragraph_format.space_before = Pt(12)
            
            logger.info("커스텀 스타일 생성 완료")
        except Exception as e:
            logger.error(f"스타일 생성 실패: {e}")
    
    def add_header_footer(self, title="FlowMate AI 보고서"):
        """헤더와 푸터 추가"""
        try:
            # 헤더 추가
            header = self.doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = title
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            header_run = header_para.runs[0]
            header_run.font.name = '맑은 고딕'
            header_run.font.size = Pt(12)
            header_run.font.color.rgb = self.theme_colors['primary']
            
            # 푸터 추가
            footer = self.doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.text = "FlowMate AI Team | "
            
            footer_run = footer_para.runs[0]
            footer_run.font.name = '맑은 고딕'
            footer_run.font.size = Pt(10)
            footer_run.font.color.rgb = self.theme_colors['light_gray']
            
            logger.info("헤더/푸터 추가 완료")
        except Exception as e:
            logger.error(f"헤더/푸터 추가 실패: {e}")
    
    def add_cover_page(self, title, subtitle="", author="FlowMate AI Team"):
        """표지 페이지 추가"""
        try:
            # 제목
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.name = '맑은 고딕'
            title_run.font.size = Pt(32)
            title_run.font.bold = True
            title_run.font.color.rgb = self.theme_colors['primary']
            
            # 공백 추가
            self.doc.add_paragraph("\n" * 3)
            
            # 부제목
            if subtitle:
                subtitle_para = self.doc.add_paragraph()
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_run = subtitle_para.add_run(subtitle)
                subtitle_run.font.name = '맑은 고딕'
                subtitle_run.font.size = Pt(18)
                subtitle_run.font.color.rgb = self.theme_colors['text']
            
            # 공백 추가
            self.doc.add_paragraph("\n" * 8)
            
            # 작성자
            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"작성자: {author}")
            author_run.font.name = '맑은 고딕'
            author_run.font.size = Pt(14)
            
            # 날짜
            date_para = self.doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
            date_run.font.name = '맑은 고딕'
            date_run.font.size = Pt(12)
            
            # 페이지 나누기
            self.doc.add_page_break()
            
            logger.info("표지 페이지 추가 완료")
        except Exception as e:
            logger.error(f"표지 페이지 추가 실패: {e}")
    
    def add_toc_placeholder(self):
        """목차 플레이스홀더 추가"""
        try:
            toc_para = self.doc.add_paragraph()
            toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_run = toc_para.add_run("목 차")
            toc_run.font.name = '맑은 고딕'
            toc_run.font.size = Pt(20)
            toc_run.font.bold = True
            toc_run.font.color.rgb = self.theme_colors['primary']
            
            self.doc.add_paragraph("\n")
            
            # 목차 안내
            toc_note = self.doc.add_paragraph()
            toc_note.add_run("※ 목차는 문서 완성 후 Word에서 '참조 > 목차' 기능을 사용하여 자동 생성하시기 바랍니다.")
            toc_note.runs[0].font.size = Pt(10)
            toc_note.runs[0].font.color.rgb = self.theme_colors['light_gray']
            
            self.doc.add_page_break()
            
            logger.info("목차 플레이스홀더 추가 완료")
        except Exception as e:
            logger.error(f"목차 추가 실패: {e}")
    
    def add_table_from_markdown(self, table_lines):
        """마크다운 테이블을 Word 테이블로 변환"""
        try:
            if not table_lines:
                return
            
            # 테이블 데이터 파싱
            table_data = []
            for line in table_lines:
                if '|' in line and not line.strip().startswith('|---'):
                    row = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if row:
                        table_data.append(row)
            
            if not table_data:
                return
            
            # 테이블 생성
            table = self.doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 테이블 데이터 입력
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table_data[0]):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = str(cell_data)
                        
                        # 헤더 행 스타일링
                        if row_idx == 0:
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            
                            # 헤더 배경색 설정
                            shading_elm = parse_xml(r'<w:shd {} w:fill="6366F1"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                            
                            # 헤더 텍스트 스타일
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.size = Pt(11)
                                    run.font.name = '맑은 고딕'
                        else:
                            # 일반 셀 스타일
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                                    run.font.name = '맑은 고딕'
            
            self.doc.add_paragraph()  # 테이블 후 여백
            logger.info("테이블 추가 완료")
        except Exception as e:
            logger.error(f"테이블 추가 실패: {e}")
    
    def process_markdown_content(self, markdown_text):
        """마크다운 내용 처리"""
        try:
            lines = markdown_text.splitlines()
            current_table_lines = []
            in_table = False
            
            for line in lines:
                line_stripped = line.strip()
                
                if not line_stripped:
                    if in_table:
                        self.add_table_from_markdown(current_table_lines)
                        current_table_lines = []
                        in_table = False
                    self.doc.add_paragraph()
                    continue
                
                # 테이블 처리
                if '|' in line_stripped and not line_stripped.startswith('#'):
                    in_table = True
                    current_table_lines.append(line_stripped)
                    continue
                elif in_table:
                    self.add_table_from_markdown(current_table_lines)
                    current_table_lines = []
                    in_table = False
                
                # 제목 처리
                if line_stripped.startswith("### "):
                    clean_line = line_stripped.replace("### ", "").strip()
                    para = self.doc.add_paragraph(clean_line, style='CustomHeading1')
                    
                elif line_stripped.startswith("#### "):
                    clean_line = line_stripped.replace("#### ", "").strip()
                    para = self.doc.add_paragraph(clean_line, style='CustomHeading2')
                    
                elif line_stripped.startswith("##### "):
                    clean_line = line_stripped.replace("##### ", "").strip()
                    para = self.doc.add_paragraph(clean_line, style='CustomHeading3')
                    
                else:
                    # 일반 문장 처리
                    para = self.doc.add_paragraph()
                    self._process_text_formatting(para, line_stripped)
            
            # 마지막 테이블 처리
            if in_table and current_table_lines:
                self.add_table_from_markdown(current_table_lines)
            
            logger.info("마크다운 내용 처리 완료")
        except Exception as e:
            logger.error(f"마크다운 처리 실패: {e}")
    
    def _process_text_formatting(self, paragraph, text):
        """텍스트 포맷팅 처리"""
        try:
            # 볼드, 이탤릭, 링크 처리
            segments = re.split(r'(\*\*[^\*]+\*\*|\*[^\*]+\*|\[[^\]]+\]\([^\)]+\))', text)
            
            for segment in segments:
                if not segment:
                    continue
                    
                run = paragraph.add_run()
                run.font.size = Pt(11)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 볼드 텍스트
                if segment.startswith("**") and segment.endswith("**"):
                    run.bold = True
                    run.text = segment[2:-2]
                    run.font.color.rgb = self.theme_colors['primary']
                    
                # 이탤릭 텍스트
                elif segment.startswith("*") and segment.endswith("*"):
                    run.italic = True
                    run.text = segment[1:-1]
                    
                # 링크
                elif re.match(r'\[[^\]]+\]\([^\)]+\)', segment):
                    link_match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', segment)
                    if link_match:
                        link_text, link_url = link_match.groups()
                        run.text = link_text
                        run.font.color.rgb = self.theme_colors['primary']
                        run.underline = True
                        
                # 일반 텍스트
                else:
                    run.text = segment
        except Exception as e:
            logger.error(f"텍스트 포맷팅 실패: {e}")
    
    def save_document(self, output_path):
        """문서 저장"""
        try:
            # 출력 디렉토리 생성
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.doc.save(output_path)
            logger.info(f"문서 저장 완료: {output_path}")
            return True
        except Exception as e:
            logger.error(f"문서 저장 실패: {e}")
            return False

def create_document_from_text(markdown_text, output_path="output/document.docx", 
                            include_cover=True, include_toc=True):
    """텍스트 내용으로부터 Word 문서 생성"""
    try:
        logger.info("문서 생성 시작")
        
        # DOCX 생성기 초기화
        generator = DOCXGenerator()
        if not generator.create_document():
            return None
        
        # 문서 제목 추출
        lines = markdown_text.splitlines()
        main_title = "FlowMate AI 보고서"
        for line in lines:
            if line.startswith("# ") or line.startswith("### "):
                main_title = line.replace("# ", "").replace("### ", "").strip()
                break
        
        # 표지 페이지 추가
        if include_cover:
            generator.add_cover_page(main_title, "AI 기반 업무 솔루션 보고서")
        
        # 목차 추가
        if include_toc:
            generator.add_toc_placeholder()
        
        # 헤더/푸터 추가
        generator.add_header_footer(main_title)
        
        # 본문 처리
        generator.process_markdown_content(markdown_text)
        
        # 문서 저장
        if generator.save_document(output_path):
            logger.info(f"문서 생성 완료: {output_path}")
            return output_path
        else:
            return None
            
    except Exception as e:
        logger.error(f"문서 생성 중 오류: {e}")
        return None

# 하위 호환성을 위한 함수
def markdown_to_styled_docx(markdown_text, output_path="output/project.docx", 
                           include_cover=True, include_toc=True):
    """기존 함수명 유지 (하위 호환성)"""
    return create_document_from_text(markdown_text, output_path, include_cover, include_toc)

if __name__ == "__main__":
    # 테스트용 샘플 마크다운
    sample_markdown = """### FlowMate AI 프로젝트 보고서

#### 1. 프로젝트 개요
**FlowMate AI**는 로컬 환경에서 구동되는 AI 기반 업무 솔루션입니다.
본 프로젝트는 *Ollama qwen2.5:7b* 모델을 활용하여 안전하고 효율적인 업무 자동화를 제공합니다.

#### 2. 주요 기능

##### 2.1 AI 채팅 어시스턴트
- **실시간 대화**: 자연스러운 한국어 대화
- **문서 기반 답변**: 업로드된 문서 내용 기반 응답
- **로컬 처리**: 외부 서버 연결 없이 안전한 처리

##### 2.2 문서 분석 기능

| 파일 형식 | 지원 여부 | 주요 기능 |
|-----------|-----------|----------|
| PDF | ✓ | 텍스트 추출, 요약 |
| DOCX | ✓ | 내용 분석, 키워드 추출 |
| CSV | ✓ | 데이터 분석, 시각화 |
| TXT | ✓ | 텍스트 처리, 요약 |

#### 3. 기술 스택

**백엔드 기술:**
- Django 웹 프레임워크
- Ollama AI 모델 (qwen2.5:7b)
- BGE-M3 임베딩
- SQLite 데이터베이스

**프론트엔드 기술:**
- HTML5, CSS3, JavaScript
- 반응형 웹 디자인
- 전문적인 보라색 테마

#### 4. 성능 지표

| 항목 | 수치 | 비고 |
|------|------|------|
| 응답 속도 | 2-5초 | 일반 질문 기준 |
| 문서 처리 | 10MB 이하 | PDF/DOCX 권장 |
| 동시 사용자 | 50명 | 로컬 서버 기준 |
| 정확도 | 85%+ | HR 예측 모델 |

#### 5. 보안 특징
- **로컬 처리**: 모든 데이터가 로컬에서 처리
- **외부 전송 없음**: 인터넷 연결 불필요
- **On-premise 배포**: 기업 내부 서버에 설치 가능

#### 6. 향후 계획
1. **모델 성능 개선**: 더 큰 모델로 업그레이드
2. **다국어 지원**: 영어, 중국어 지원 추가
3. **API 확장**: RESTful API 제공
4. **모바일 앱**: 네이티브 앱 개발

---

본 보고서는 FlowMate AI 프로젝트의 현재 상태와 향후 발전 방향을 종합적으로 정리한 문서입니다.
더 자세한 정보가 필요하시면 [FlowMate AI Team](mailto:info@flowmate.ai)으로 연락 주시기 바랍니다.
"""
    
    # 문서 생성 테스트
    output_path = create_document_from_text(sample_markdown, "output/test_document.docx")
    
    if output_path:
        print(f"✅ 문서 생성 성공: {output_path}")
        print("주요 개선사항:")
        print("- 한국어 전용 처리")
        print("- 예외 처리 강화")
        print("- 로깅 시스템 추가")
        print("- 클래스 기반 구조")
        print("- 전문적인 표지 및 목차")
        print("- 테이블 자동 변환")
        print("- 보라색 테마 적용")
    else:
        print("❌ 문서 생성 실패")